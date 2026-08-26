"""Generate a clean, professional Word document (.docx) for Approach.docx."""

import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


def set_cell_background(cell, fill_hex):
    """Set background color of a table cell."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tc_pr.append(shd)


def create_approach_document(output_path: str):
    doc = docx.Document()

    # Page setup
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # Style definitions
    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Calibri'
    normal_style.font.size = Pt(11)
    normal_style.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    t_run = title.add_run("Engineering Approach & System Evolution")
    t_run.font.name = 'Calibri'
    t_run.font.size = Pt(24)
    t_run.font.bold = True
    t_run.font.color.rgb = RGBColor(0x1B, 0x36, 0x5D)

    subtitle = doc.add_paragraph()
    s_run = subtitle.add_run("Video Dialogue Retrieval and Precision Frame Localization")
    s_run.font.name = 'Calibri'
    s_run.font.size = Pt(13)
    s_run.font.italic = True
    s_run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    doc.add_paragraph()  # spacing

    # 1. Initial Thoughts & Mental Models
    h1 = doc.add_heading("1. Initial Thoughts & Brainstorming", level=1)
    h1.style.font.color.rgb = RGBColor(0x1B, 0x36, 0x5D)

    doc.add_paragraph(
        "After carefully reading the problem statement, the core goal was defined: Given a video URL and a target dialogue "
        "phrase (e.g., \"My mind rebels at stagnation\"), locate the exact moment in the video where the dialogue is spoken "
        "and retrieve the corresponding visual frame."
    )
    doc.add_paragraph(
        "The initial brainstorming session focused on understanding video as a multi-modal time-series data structure containing "
        "two synchronized components: visual images (frames) and an audio soundtrack (speech). Several early ideas were noted:"
    )

    p = doc.add_paragraph(style='List Bullet')
    p.add_run("1-to-1 Frame Slicing Hypothesis: ").bold = True
    p.add_run("Initially thought of dividing the video into frames and slicing the audio into matching chunks for each video frame (e.g. 1/25th of a second).")

    p = doc.add_paragraph(style='List Bullet')
    p.add_run("Speech Activity Isolation: ").bold = True
    p.add_run("Recognized that non-speech sections (silence, background music) should be pruned to shrink the search space before running speech recognition.")

    p = doc.add_paragraph(style='List Bullet')
    p.add_run("The Stopword Dilution Problem: ").bold = True
    p.add_run("Realized that spoken dialogue spans multiple frames, and common stopwords like \"the\" appear across hundreds of frames. Searching for common words would yield useless false positives.")

    p = doc.add_paragraph(style='List Bullet')
    p.add_run("Keyword Anchoring & TF-IDF Rarity: ").bold = True
    p.add_run("Formulated the concept of using term rarity (TF-IDF) to find the most unique, informative word in the target sentence (e.g. \"rebels\" or \"stagnation\") and searching only around occurrences of that anchor.")

    # 2. Review and Critical Architectural Fixes
    h1 = doc.add_heading("2. Review & Critical Architectural Fixes", level=1)
    h1.style.font.color.rgb = RGBColor(0x1B, 0x36, 0x5D)

    doc.add_paragraph(
        "When reviewing these initial ideas with Claude, critical feedback helped pivot away from naive assumptions toward "
        "a mathematically sound and scalable architecture:"
    )

    p = doc.add_paragraph()
    p_run = p.add_run("The Main Fix: Stop Thinking in Discrete Audio Frames")
    p_run.bold = True
    p_run.font.color.rgb = RGBColor(0x00, 0x5B, 0x94)

    doc.add_paragraph(
        "Attempting a 1:1 mapping between video frames and audio frames contradicts how Automatic Speech Recognition (ASR) works. "
        "Human speech phonemes and words span 200ms to 1.5 seconds. Modern ASR architectures (such as Whisper) take a continuous audio stream "
        "and output text with continuous timestamps in seconds (float), not discrete frame numbers. Once the continuous timestamp (t) "
        "is found, converting to a video frame index is simply a mathematical calculation: Frame = round(t * FPS)."
    )

    doc.add_paragraph(
        "Additional key corrections included:\n"
        "• Audio-First Acquisition: Downloading heavy 1080p/4K video files upfront is wasteful. Downloading lightweight 16kHz mono audio "
        "first allows fast transcription and search, deferring full video download until a match is verified.\n"
        "• Acoustic Fingerprinting: Video URLs and container metadata shift across platforms. Chromaprint acoustic fingerprints uniquely identify "
        "audio content regardless of re-encoding or container format, enabling robust SQLite caching."
    )

    # 3. The 7-Step Modular Framework
    h1 = doc.add_heading("3. The 7-Step Modular Framework", level=1)
    h1.style.font.color.rgb = RGBColor(0x1B, 0x36, 0x5D)

    doc.add_paragraph("To keep development organized and verifiable, the end-to-end pipeline was decomposed into seven modular steps:")

    steps = [
        ("Step 1: Video Acquisition Strategy", "Support remote URLs (yt-dlp) and local files with lightweight metadata extraction (FPS, duration, resolution) without full upfront video downloads."),
        ("Step 2: Audio Extraction", "Extract standardized single-channel 16kHz PCM WAV audio for optimal acoustic fingerprinting and ASR processing."),
        ("Step 3: Speech Separation & VAD", "Apply Voice Activity Detection (VAD) to filter out non-speech silence and music, minimizing transcription latency."),
        ("Step 4: Speech-to-Text Conversion (ASR)", "Run neural speech recognition (Faster-Whisper) with word-level alignment, generating precise word start and end timestamps."),
        ("Step 5: Efficient Dialogue Search", "Search the transcript using Exact Phrase Matching, Fuzzy Sliding Window, and Rare-Anchor Search with Inverted Indexing."),
        ("Step 6: Frame Localization & Retrieval", "Calculate discrete frame indices from match timestamps (Frame = round(t * FPS)) and extract crisp JPEG frames via FFmpeg input seeking."),
        ("Step 7: Structured Output & Reporting", "Generate structured JSON reports containing match ranks, confidence scores, visual frame image paths, and timing profiles.")
    ]

    for title_text, desc in steps:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f"{title_text}: ").bold = True
        p.add_run(desc)

    # 4. Prototyping Strategy: Kaggle T4 GPU
    h1 = doc.add_heading("4. Prototyping Strategy (Kaggle T4 GPU & Notebooks)", level=1)
    h1.style.font.color.rgb = RGBColor(0x1B, 0x36, 0x5D)

    doc.add_paragraph(
        "Because local hardware lacked a dedicated NVIDIA GPU, running Whisper models (especially small and medium sizes) locally was slow for rapid iteration. "
        "The workflow was initially prototyped in Jupyter Notebooks on Kaggle using free NVIDIA T4 GPU instances. This allowed rapid testing of ASR models, "
        "interactive inspection of word timestamps, and quick validation of search algorithms on real 50+ minute video transcripts."
    )

    # 5. Evolution Across Notebook Versions
    h1 = doc.add_heading("5. Iterative Refinements Across Notebook Versions (v1 -> v2 -> v3)", level=1)
    h1.style.font.color.rgb = RGBColor(0x1B, 0x36, 0x5D)

    doc.add_paragraph("The notebook evolved through three iterations in the Notebooks/ folder:")

    # Table for versions
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Version"
    hdr_cells[1].text = "Key Changes & Enhancements"
    hdr_cells[2].text = "Identified Bottlenecks & Fixes"

    for cell in hdr_cells:
        set_cell_background(cell, "1B365D")
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.bold = True
                r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    version_data = [
        ("v1 (Initial Prototype)", 
         "• Linear implementation of the 7 steps on Kaggle.\n• Manual Silero VAD + per-segment FFmpeg slicing.\n• Basic exact and difflib fuzzy search.",
         "• Boundary Overlap Bug: Padding VAD chunks caused adjacent segments to overlap, transcribing boundary words twice.\n• Stale IDF Bug: Term frequencies were global, leaking stats across videos.\n• Hardcoded model size ('small')."),
        ("v2 (Streamlined VAD & Benchmarking)",
         "• Switched to faster-whisper's native vad_filter=True (single-pass, zero disk slicing).\n• Dynamic per-call IDF rarity calculation.\n• Added multi-model comparison harness.\n• Added sentence-transformers embedding scorer.",
         "• Cache Artifact: Model benchmark reused cached transcripts, making 'small' look instantaneous (0.0038s).\n• Large video files still downloaded upfront.\n• Pure Python difflib search was relatively slow."),
        ("v3 (Audio-First & Inverted Index)",
         "• Audio-First Lazy Video Acquisition (downloads audio only; defers full video until match is verified).\n• Chromaprint Acoustic Fingerprinting (AcoustID) + SQLite dedup.\n• Inverted Index (word -> positions) for O(1) anchor lookups.\n• RapidFuzz C++ scoring (10x faster than difflib).",
         "• Formed the robust, complete algorithmic foundation ready for conversion into a modular production Python project.")
    ]

    for v_num, changes, fixes in version_data:
        row_cells = table.add_row().cells
        row_cells[0].text = v_num
        row_cells[1].text = changes
        row_cells[2].text = fixes
        set_cell_background(row_cells[0], "F0F4F8")

    doc.add_paragraph()

    # 6. Conversion to Industry-Standard Python Package
    h1 = doc.add_heading("6. Transition to Production Python Package by Antigravity", level=1)
    h1.style.font.color.rgb = RGBColor(0x1B, 0x36, 0x5D)

    doc.add_paragraph(
        "Once the algorithmic workflow was verified in v3, the project was restructured by Antigravity into a modular, production-ready "
        "Python package (`video_dialogue_retrieval/`) following software engineering best practices:"
    )

    p = doc.add_paragraph(style='List Bullet')
    p.add_run("Decoupled Architecture: ").bold = True
    p.add_run("Separated into independent layers: config, core models, database (SQLite), audio downloader, acoustic fingerprinting, ASR transcriber, search engines, video frame extractor, and pipeline orchestrator.")

    p = doc.add_paragraph(style='List Bullet')
    p.add_run("Resilient Audio Fingerprinting: ").bold = True
    p.add_run("Implemented dual-engine Chromaprint extraction (pyacoustid with automatic fallback to FFmpeg's built-in chromaprint muxer) for seamless offline operation without requiring external fpcalc binaries.")

    p = doc.add_paragraph(style='List Bullet')
    p.add_run("Multi-Anchor Retry & Sliding Window Fallback: ").bold = True
    p.add_run("Enhanced rare-anchor search so that if the top anchor word is missing due to ASR misspelling, it retries with secondary candidate anchors before automatically falling back to full sliding window search.")

    p = doc.add_paragraph(style='List Bullet')
    p.add_run("Duration Verification on SQLite Dedup: ").bold = True
    p.add_run("Compares input video duration against SQLite records to ensure that trimmed re-uploads with differing frame counts are not erroneously served cached transcripts.")

    p = doc.add_paragraph(style='List Bullet')
    p.add_run("Optimized Video Downloads: ").bold = True
    p.add_run("Configured yt-dlp format selection to download lower resolution video streams (height <= 720p/480p) during lazy acquisition for high-speed frame extraction.")

    p = doc.add_paragraph(style='List Bullet')
    p.add_run("Automated Test Suite: ").bold = True
    p.add_run("Built a comprehensive 34-test pytest suite covering SQLite persistence, fingerprinting determinism, text normalization, similarity scorers, inverted indexing, search dispatch, and end-to-end pipeline execution.")

    p = doc.add_paragraph(style='List Bullet')
    p.add_run("Rich Command-Line Interface (CLI): ").bold = True
    p.add_run("Provides full CLI commands: search, benchmark, inspect, and clear-cache with progress indicators and structured tabular outputs.")

    # 7. Algorithmic Parameter Clarification
    h1 = doc.add_heading("7. Parameter & Configuration Reference", level=1)
    h1.style.font.color.rgb = RGBColor(0x1B, 0x36, 0x5D)

    doc.add_paragraph(
        "• Whisper Decoding Beam Size: Default beam size is 5 in Whisper ASR transcription (configurable in PipelineConfig).\n"
        "• Dialogue Search Window Tolerances: Length tolerance is 2 (evaluates window lengths [N-2, N+2]), and context expansion window is 2 (expands context around anchor positions by ±2 words).\n"
        "• Audio Sample Rate: 16,000 Hz single-channel mono PCM WAV.\n"
        "• VAD Settings: vad_filter=True with min_silence_duration_ms=300."
    )

    doc.save(output_path)
    print(f"Successfully generated Word document: {output_path}")


if __name__ == "__main__":
    create_approach_document("c:/Users/computer/Desktop/Quest1/Approach.docx")
